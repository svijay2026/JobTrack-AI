import io
import pytest
from docx import Document
from fastapi.testclient import TestClient


def create_auth_user(client: TestClient, email: str = "jobuser@example.com", password: str = "SecurePass123!"):
    """Helper to register and login user, returning authorization headers."""
    client.post(
        "/api/v1/auth/register",
        json={"email": email, "password": password, "full_name": "Job Candidate"},
    )
    login_resp = client.post(
        "/api/v1/auth/login",
        data={"username": email, "password": password},
    )
    token = login_resp.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def create_dummy_resume(client: TestClient, headers: dict) -> int:
    """Uploads a dummy resume and returns its ID."""
    doc = Document()
    doc.add_heading("Resume", level=1)
    doc.add_paragraph("Python, FastAPI developer")
    buf = io.BytesIO()
    doc.save(buf)

    resp = client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("test_resume.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    return resp.json()["id"]


def test_create_and_get_job_application(client: TestClient):
    headers = create_auth_user(client, email="candidate_create@example.com")
    
    payload = {
        "company_name": "Google",
        "job_title": "Senior Python Backend Engineer",
        "job_location": "Mountain View, CA / Remote",
        "job_type": "full_time",
        "salary_range": "$160,000 - $210,000",
        "job_description": "We are seeking a senior engineer to design distributed systems using Python and Kubernetes.",
        "job_url": "https://careers.google.com/jobs/results/12345",
        "status": "applied",
        "notes": "Spoke to recruiter Sarah on LinkedIn.",
    }

    response = client.post("/api/v1/jobs/", json=payload, headers=headers)
    assert response.status_code == 201
    job_data = response.json()
    assert job_data["company_name"] == "Google"
    assert job_data["job_title"] == "Senior Python Backend Engineer"
    assert job_data["status"] == "applied"
    job_id = job_data["id"]

    # Get by ID
    get_resp = client.get(f"/api/v1/jobs/{job_id}", headers=headers)
    assert get_resp.status_code == 200
    assert get_resp.json()["id"] == job_id


def test_create_job_with_linked_resume(client: TestClient):
    headers = create_auth_user(client, email="resume_linker@example.com")
    resume_id = create_dummy_resume(client, headers)

    payload = {
        "company_name": "Microsoft",
        "job_title": "Cloud Architect",
        "status": "applied",
        "resume_id": resume_id,
    }

    response = client.post("/api/v1/jobs/", json=payload, headers=headers)
    assert response.status_code == 201
    assert response.json()["resume_id"] == resume_id

    # Test invalid resume ID -> 404
    bad_payload = {
        "company_name": "Invalid Corp",
        "job_title": "Dev",
        "resume_id": 99999,
    }
    bad_resp = client.post("/api/v1/jobs/", json=bad_payload, headers=headers)
    assert bad_resp.status_code == 404


def test_list_and_filter_jobs(client: TestClient):
    headers = create_auth_user(client, email="filter_user@example.com")

    # Create 3 jobs with different statuses
    client.post("/api/v1/jobs/", json={"company_name": "Amazon", "job_title": "SDE II", "status": "applied"}, headers=headers)
    client.post("/api/v1/jobs/", json={"company_name": "Meta", "job_title": "Production Engineer", "status": "interviewing"}, headers=headers)
    client.post("/api/v1/jobs/", json={"company_name": "Netflix", "job_title": "Staff Engineer", "status": "offered"}, headers=headers)

    # List all
    all_resp = client.get("/api/v1/jobs/", headers=headers)
    assert all_resp.status_code == 200
    assert len(all_resp.json()) == 3

    # Filter by interviewing
    interview_resp = client.get("/api/v1/jobs/?status=interviewing", headers=headers)
    assert interview_resp.status_code == 200
    assert len(interview_resp.json()) == 1
    assert interview_resp.json()[0]["company_name"] == "Meta"


def test_search_jobs(client: TestClient):
    headers = create_auth_user(client, email="search_user@example.com")

    client.post("/api/v1/jobs/", json={"company_name": "Stripe", "job_title": "Fintech Platform Engineer"}, headers=headers)
    client.post("/api/v1/jobs/", json={"company_name": "Spotify", "job_title": "Audio Pipeline Developer"}, headers=headers)

    # Search by company
    stripe_search = client.get("/api/v1/jobs/?search=Stripe", headers=headers)
    assert len(stripe_search.json()) == 1
    assert stripe_search.json()[0]["company_name"] == "Stripe"

    # Search by job title keyword
    audio_search = client.get("/api/v1/jobs/?search=Audio", headers=headers)
    assert len(audio_search.json()) == 1
    assert audio_search.json()[0]["company_name"] == "Spotify"


def test_patch_status_kanban(client: TestClient):
    headers = create_auth_user(client, email="kanban_user@example.com")

    create_resp = client.post(
        "/api/v1/jobs/",
        json={"company_name": "Uber", "job_title": "Backend Tech Lead", "status": "applied"},
        headers=headers,
    )
    job_id = create_resp.json()["id"]

    # Transition to interviewing
    patch_resp = client.patch(f"/api/v1/jobs/{job_id}/status", json={"status": "interviewing"}, headers=headers)
    assert patch_resp.status_code == 200
    assert patch_resp.json()["status"] == "interviewing"

    # Transition to offered
    patch_resp2 = client.patch(f"/api/v1/jobs/{job_id}/status", json={"status": "offered"}, headers=headers)
    assert patch_resp2.status_code == 200
    assert patch_resp2.json()["status"] == "offered"


def test_job_stats_funnel(client: TestClient):
    headers = create_auth_user(client, email="stats_user@example.com")

    client.post("/api/v1/jobs/", json={"company_name": "Co 1", "job_title": "Dev", "status": "wishlist"}, headers=headers)
    client.post("/api/v1/jobs/", json={"company_name": "Co 2", "job_title": "Dev", "status": "applied"}, headers=headers)
    client.post("/api/v1/jobs/", json={"company_name": "Co 3", "job_title": "Dev", "status": "interviewing"}, headers=headers)
    client.post("/api/v1/jobs/", json={"company_name": "Co 4", "job_title": "Dev", "status": "offered"}, headers=headers)
    client.post("/api/v1/jobs/", json={"company_name": "Co 5", "job_title": "Dev", "status": "rejected"}, headers=headers)

    stats_resp = client.get("/api/v1/jobs/stats", headers=headers)
    assert stats_resp.status_code == 200
    data = stats_resp.json()
    assert data["total_applications"] == 5
    assert data["wishlist"] == 1
    assert data["applied"] == 1
    assert data["interviewing"] == 1
    assert data["offered"] == 1
    assert data["rejected"] == 1
    assert data["interview_rate_percent"] > 0
    assert data["offer_rate_percent"] > 0


def test_update_and_delete_job(client: TestClient):
    headers = create_auth_user(client, email="update_delete_user@example.com")

    create_resp = client.post(
        "/api/v1/jobs/",
        json={"company_name": "Airbnb", "job_title": "Full Stack Engineer", "status": "applied"},
        headers=headers,
    )
    job_id = create_resp.json()["id"]

    # Update job
    put_resp = client.put(
        f"/api/v1/jobs/{job_id}",
        json={"company_name": "Airbnb Inc.", "salary_range": "$180,000", "notes": "Final round completed"},
        headers=headers,
    )
    assert put_resp.status_code == 200
    assert put_resp.json()["company_name"] == "Airbnb Inc."
    assert put_resp.json()["salary_range"] == "$180,000"

    # Delete job
    del_resp = client.delete(f"/api/v1/jobs/{job_id}", headers=headers)
    assert del_resp.status_code == 200

    # Verify 404
    get_resp = client.get(f"/api/v1/jobs/{job_id}", headers=headers)
    assert get_resp.status_code == 404


def test_jobs_cross_user_isolation(client: TestClient):
    user1_headers = create_auth_user(client, email="job_owner1@example.com")
    user2_headers = create_auth_user(client, email="job_owner2@example.com")

    create_resp = client.post(
        "/api/v1/jobs/",
        json={"company_name": "OpenAI", "job_title": "Research Engineer"},
        headers=user1_headers,
    )
    job_id = create_resp.json()["id"]

    # User 2 attempts to get User 1's job
    assert client.get(f"/api/v1/jobs/{job_id}", headers=user2_headers).status_code == 404

    # User 2 attempts to patch User 1's job
    assert client.patch(f"/api/v1/jobs/{job_id}/status", json={"status": "offered"}, headers=user2_headers).status_code == 404

    # User 2 attempts to delete User 1's job
    assert client.delete(f"/api/v1/jobs/{job_id}", headers=user2_headers).status_code == 404
