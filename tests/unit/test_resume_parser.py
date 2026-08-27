import io
import pytest
from docx import Document
from pypdf import PdfWriter

from app.services.resume_parser import resume_parser


def test_clean_text():
    raw = "  Hello   world!  \r\n\r\n\r\nThis is a   test.  \t "
    cleaned = resume_parser.clean_text(raw)
    assert cleaned == "Hello world!\n\nThis is a test."


def test_extract_contact_info():
    sample_text = """
    John Doe
    Email: john.doe@example.com
    Phone: (555) 123-4567
    LinkedIn: linkedin.com/in/johndoe
    GitHub: github.com/johndoe-dev
    """
    contact = resume_parser.extract_contact_info(sample_text)
    assert contact["email"] == "john.doe@example.com"
    assert contact["phone"] == "(555) 123-4567"
    assert "johndoe" in contact["linkedin"]
    assert "johndoe-dev" in contact["github"]


def test_extract_skills():
    sample_text = """
    Experienced Software Engineer skilled in Python, FastAPI, Docker, and PostgreSQL.
    Also familiar with React, Next.js, and CI/CD pipelines using GitHub Actions.
    Worked with C++ and Go microservices.
    """
    skills = resume_parser.extract_skills(sample_text)
    assert "Python" in skills
    assert "FastAPI" in skills
    assert "Docker" in skills
    assert "PostgreSQL" in skills
    assert "React" in skills
    assert "Next.js" in skills
    assert "C++" in skills
    assert "Go" in skills
    assert "GitHub Actions" in skills


def test_extract_education():
    sample_text = """
    Education:
    Bachelor of Science in Computer Science, 2021
    Master of Science in Artificial Intelligence, 2023
    """
    education = resume_parser.extract_education(sample_text)
    assert len(education) >= 2
    degrees = [item["degree"].lower() for item in education]
    assert any("bachelor" in d for d in degrees)
    assert any("master" in d for d in degrees)


def test_extract_experience_years_explicit():
    sample_text = "Senior Python Developer with 5+ years of experience building APIs."
    years = resume_parser.extract_experience_years(sample_text)
    assert years == 5.0


def test_extract_experience_years_date_ranges():
    sample_text = """
    Software Engineer (2018 - 2022) - Built backend microservices.
    Senior Engineer (2022 - Present) - Leading architecture design.
    """
    years = resume_parser.extract_experience_years(sample_text)
    assert years >= 5.0


def test_parse_docx():
    # Create in-memory DOCX document
    doc = Document()
    doc.add_heading("Jane Doe - Resume", level=1)
    doc.add_paragraph("Email: jane.doe@domain.com | Phone: 123-456-7890")
    doc.add_paragraph("Skills: Python, FastAPI, MySQL, Redis, Docker, Kubernetes")
    doc.add_paragraph("Bachelor of Technology in Computer Science, 2020")
    doc.add_paragraph("Experience: 2020 - 2024 Software Developer at Acme Inc.")

    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    result = resume_parser.parse(file_bytes=docx_bytes, file_extension=".docx")
    assert "Jane Doe" in result["parsed_text"]
    assert result["contact_info"]["email"] == "jane.doe@domain.com"
    assert "Python" in result["skills"]
    assert "FastAPI" in result["skills"]
    assert "Docker" in result["skills"]
    assert result["experience_years"] >= 4.0
