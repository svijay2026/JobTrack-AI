import io
import re
from datetime import datetime
from typing import Any, Dict, List, Set, Tuple
from pypdf import PdfReader
from docx import Document


# Comprehensive taxonomy of skills for technology, engineering, data & management
SKILL_TAXONOMY: Dict[str, List[str]] = {
    "Programming Languages": [
        "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "C", "Go", "Golang",
        "Rust", "Ruby", "PHP", "Swift", "Kotlin", "Scala", "R", "Dart", "MATLAB", "Perl", "Bash", "Shell"
    ],
    "Web & Backend Frameworks": [
        "FastAPI", "Django", "Flask", "Node.js", "Express.js", "Express", "NestJS", "Spring", "Spring Boot",
        "ASP.NET", ".NET Core", "Ruby on Rails", "Laravel", "Gin", "Echo", "GraphQL", "REST", "RESTful API",
        "gRPC", "Microservices"
    ],
    "Frontend & Mobile": [
        "React", "React.js", "Next.js", "Vue", "Vue.js", "Nuxt.js", "Angular", "Svelte", "Redux", "Zustand",
        "Tailwind CSS", "Bootstrap", "HTML5", "CSS3", "SASS", "SCSS", "Flutter", "React Native", "Android", "iOS"
    ],
    "Databases & Caching": [
        "SQL", "MySQL", "PostgreSQL", "Postgres", "SQLite", "MongoDB", "Redis", "Elasticsearch", "Cassandra",
        "DynamoDB", "Oracle", "SQL Server", "MSSQL", "MariaDB", "Neo4j", "Supabase", "Firebase"
    ],
    "Cloud & DevOps": [
        "AWS", "Amazon Web Services", "Azure", "GCP", "Google Cloud", "Docker", "Kubernetes", "K8s",
        "Terraform", "Ansible", "Jenkins", "GitHub Actions", "GitLab CI", "CI/CD", "Linux", "Nginx",
        "Prometheus", "Grafana", "Cloudflare", "Serverless"
    ],
    "AI, ML & Data Science": [
        "Machine Learning", "Deep Learning", "Artificial Intelligence", "NLP", "Natural Language Processing",
        "Computer Vision", "LLM", "Generative AI", "PyTorch", "TensorFlow", "Keras", "Scikit-Learn",
        "Pandas", "NumPy", "OpenCV", "Hugging Face", "LangChain", "LlamaIndex", "BERT", "GPT", "OpenAI"
    ],
    "Tools & Practices": [
        "Git", "GitHub", "GitLab", "Bitbucket", "Jira", "Confluence", "Agile", "Scrum", "TDD", "Unit Testing",
        "Pytest", "Jest", "Postman", "Swagger", "Docker Compose", "Webpack", "Vite"
    ]
}


class ResumeParser:
    """
    Intelligent extraction engine for resumes in PDF and DOCX formats.
    Extracts raw text, contact information, education, experience, and domain skills.
    """

    def extract_text_from_pdf(self, file_bytes: bytes) -> str:
        """Extracts text from PDF bytes using pypdf."""
        text_content = []
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
        except Exception as e:
            raise ValueError(f"Failed to parse PDF document: {str(e)}")
        
        return "\n".join(text_content)

    def extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extracts text from DOCX bytes using python-docx."""
        text_content = []
        try:
            doc = Document(io.BytesIO(file_bytes))
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_content.append(" | ".join(row_text))
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX document: {str(e)}")
        
        return "\n".join(text_content)

    def clean_text(self, text: str) -> str:
        """Normalizes whitespaces, removes unwanted control characters and trims lines."""
        if not text:
            return ""
        text = re.sub(r"\r\n|\r", "\n", text)
        lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
        
        cleaned_lines: List[str] = []
        for line in lines:
            if line:
                cleaned_lines.append(line)
            elif cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
                
        return "\n".join(cleaned_lines).strip()

    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extracts email, phone, LinkedIn, GitHub, and portfolio links from text."""
        contact: Dict[str, Any] = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "portfolio": None,
        }

        # Email Extraction
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        email_match = re.search(email_pattern, text)
        if email_match:
            contact["email"] = email_match.group(0)

        # Phone Number Extraction (supports international + standard formats)
        phone_pattern = r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}"
        phone_match = re.search(phone_pattern, text)
        if phone_match:
            contact["phone"] = phone_match.group(0).strip()

        # LinkedIn Extraction
        linkedin_pattern = r"(?:https?://)?(?:www\.)?linkedin\.com/in/([A-Za-z0-9_-]+)"
        linkedin_match = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_match:
            contact["linkedin"] = f"https://linkedin.com/in/{linkedin_match.group(1)}"

        # GitHub Extraction
        github_pattern = r"(?:https?://)?(?:www\.)?github\.com/([A-Za-z0-9_-]+)"
        github_match = re.search(github_pattern, text, re.IGNORECASE)
        if github_match:
            contact["github"] = f"https://github.com/{github_match.group(1)}"

        return contact

    def extract_skills(self, text: str) -> List[str]:
        """
        Matches technical and professional skills using boundary-safe regex against the skill catalog.
        Handles standard alphanumeric skills as well as special characters (C++, C#, .NET, Node.js).
        """
        found_skills: Set[str] = set()
        lower_text = f" {text.lower()} "

        for category, skills in SKILL_TAXONOMY.items():
            for skill in skills:
                lower_skill = skill.lower()
                escaped_skill = re.escape(lower_skill)

                if "++" in lower_skill:
                    pattern = rf"(?<![a-zA-Z0-9]){escaped_skill}(?![a-zA-Z0-9\+])"
                elif "#" in lower_skill:
                    pattern = rf"(?<![a-zA-Z0-9]){escaped_skill}(?![a-zA-Z0-9#])"
                elif lower_skill.startswith("."):
                    pattern = rf"(?<![a-zA-Z0-9]){escaped_skill}\b"
                else:
                    pattern = rf"\b{escaped_skill}\b"

                if re.search(pattern, lower_text):
                    found_skills.add(skill)

        # Sort alphabetically for consistency
        return sorted(list(found_skills))

    def extract_education(self, text: str) -> List[Dict[str, Any]]:
        """Extracts degrees, institutions, and graduation years using regex heuristics."""
        education_list = []
        degree_patterns = [
            r"\b(Bachelor(?:'s)?(?:\s+of\s+[A-Za-z]+)?|B\.?S\.?|B\.?A\.?|B\.?Tech\.?|B\.?E\.?)\b",
            r"\b(Master(?:'s)?(?:\s+of\s+[A-Za-z]+)?|M\.?S\.?|M\.?A\.?|M\.?Tech\.?|M\.?B\.?A\.?)\b",
            r"\b(Ph\.?D\.?|Doctorate|Doctor of Philosophy)\b",
            r"\b(Associate(?:'s)?(?:\s+Degree)?)\b",
        ]

        lines = text.split("\n")
        for i, line in enumerate(lines):
            for pattern in degree_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    degree_found = match.group(0).strip()
                    # Check for year in same line or next line
                    year_match = re.search(r"\b(19\d{2}|20\d{2})\b", line)
                    if not year_match and i + 1 < len(lines):
                        year_match = re.search(r"\b(19\d{2}|20\d{2})\b", lines[i + 1])
                    
                    year = year_match.group(0) if year_match else None

                    education_list.append({
                        "degree": degree_found,
                        "raw_entry": line.strip(),
                        "year": year
                    })
                    break

        return education_list

    def extract_experience_years(self, text: str) -> float:
        """
        Estimates years of experience based on explicit experience mentions
        and employment date range spans (e.g. 2018 - 2023).
        """
        # Check explicit mentions e.g., "5+ years of experience", "3 years of software engineering"
        explicit_pattern = r"(\d+(?:\.\d+)?)\+?\s*(?:years|yrs)\s+(?:of\s+)?experience"
        explicit_match = re.search(explicit_pattern, text, re.IGNORECASE)
        if explicit_match:
            try:
                return float(explicit_match.group(1))
            except ValueError:
                pass

        # Check date range patterns (e.g. 2018 - Present, 2019 - 2022)
        current_year = datetime.now().year
        year_range_pattern = r"\b(20\d{2}|19\d{2})\s*(?:-|–|to)\s*(20\d{2}|19\d{2}|present|current)\b"
        matches = re.findall(year_range_pattern, text, re.IGNORECASE)

        total_years = 0.0
        seen_ranges: Set[Tuple[int, int]] = set()

        for start_str, end_str in matches:
            try:
                start_year = int(start_str)
                if end_str.lower() in ("present", "current"):
                    end_year = current_year
                else:
                    end_year = int(end_str)

                if start_year <= end_year and 1980 <= start_year <= current_year:
                    range_tuple = (start_year, end_year)
                    if range_tuple not in seen_ranges:
                        seen_ranges.add(range_tuple)
                        diff = end_year - start_year
                        total_years += max(1.0, float(diff))
            except ValueError:
                continue

        # Cap estimated total years at reasonable bound
        return min(round(total_years, 1), 40.0)

    def parse(self, file_bytes: bytes, file_extension: str) -> Dict[str, Any]:
        """
        Orchestrates full parsing pipeline for a given resume file.
        Returns clean text, skills list, estimated experience, education list, and contact info.
        """
        ext = file_extension.lower()
        if ext == ".pdf":
            raw_text = self.extract_text_from_pdf(file_bytes)
        elif ext in [".docx", ".doc"]:
            raw_text = self.extract_text_from_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

        cleaned_text = self.clean_text(raw_text)
        skills = self.extract_skills(cleaned_text)
        contact_info = self.extract_contact_info(cleaned_text)
        education = self.extract_education(cleaned_text)
        experience_years = self.extract_experience_years(cleaned_text)

        return {
            "parsed_text": cleaned_text,
            "skills": skills,
            "experience_years": experience_years,
            "education": education,
            "contact_info": contact_info,
        }


resume_parser = ResumeParser()
