import io
import pytest
from docx import Document
from fastapi.testclient import TestClient


def create_auth_user(client: TestClient, email: str = "testuser@example.com", password: str = "SecurePass123!"):
    """Helper to register and authenticate a test user, returning auth headers."""
    # Register
    client.post(
        "/api/v1/auth/register",
        json={"email": email, "password": password, "full_name": "Test Candidate"},
    )
    # Login
    login_resp = client.post(
        "/api/v1/auth/login",
        data={"username": email, "password": password},
    )
    token = login_resp.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def create_sample_docx_bytes():
    """Generates a valid test DOCX file in memory."""
    doc = Document()
    doc.add_heading("Alex Rivera", level=1)
    doc.add_paragraph("Email: alex.rivera@techjob.io | Phone: +1 555-019-2831")
    doc.add_paragraph("LinkedIn: https://linkedin.com/in/alexrivera | GitHub: https://github.com/alexrivera")
    doc.add_paragraph("Skills: Python, FastAPI, Docker, Kubernetes, AWS, PostgreSQL, React, TypeScript")
    doc.add_paragraph("Education: Bachelor of Science in Computer Science, 2020")
    doc.add_paragraph("Experience: 2020 - 2024 Software Engineer at Global Tech")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_upload_and_parse_resume_api(client: TestClient):
    headers = create_auth_user(client, email="resume_tester@example.com")
    docx_bytes = create_sample_docx_bytes()

    response = client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("alex_rivera_resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"is_primary": "true"},
    )

    assert response.status_code == 201
    data = response.json()
    assert data["file_name"] == "alex_rivera_resume.docx"
    assert data["is_primary"] is True
    assert "Python" in data["skills"]
    assert "FastAPI" in data["skills"]
    assert "Docker" in data["skills"]
    assert data["contact_info"]["email"] == "alex.rivera@techjob.io"
    assert data["experience_years"] >= 4.0
    assert len(data["education"]) > 0


def test_upload_invalid_file_extension(client: TestClient):
    headers = create_auth_user(client, email="bad_file_user@example.com")
    
    response = client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("malicious_script.sh", b"echo 'hello'", "application/x-sh")},
    )

    assert response.status_code == 400
    assert "Unsupported file extension" in response.json()["detail"]


def test_list_resumes_api(client: TestClient):
    headers = create_auth_user(client, email="multi_resume_user@example.com")
    docx_bytes = create_sample_docx_bytes()

    # Upload first resume
    client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("resume_v1.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    # Upload second resume
    client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("resume_v2.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    response = client.get("/api/v1/resumes/", headers=headers)
    assert response.status_code == 200
    resumes = response.json()
    assert len(resumes) == 2


def test_get_resume_by_id_and_primary(client: TestClient):
    headers = create_auth_user(client, email="primary_test_user@example.com")
    docx_bytes = create_sample_docx_bytes()

    upload_resp = client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("primary_resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    resume_id = upload_resp.json()["id"]

    # Get by ID
    get_resp = client.get(f"/api/v1/resumes/{resume_id}", headers=headers)
    assert get_resp.status_code == 200
    assert get_resp.json()["id"] == resume_id

    # Get primary
    primary_resp = client.get("/api/v1/resumes/primary", headers=headers)
    assert primary_resp.status_code == 200
    assert primary_resp.json()["id"] == resume_id


def test_download_and_delete_resume(client: TestClient):
    headers = create_auth_user(client, email="download_user@example.com")
    docx_bytes = create_sample_docx_bytes()

    upload_resp = client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("download_me.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    resume_id = upload_resp.json()["id"]

    # Download file
    download_resp = client.get(f"/api/v1/resumes/{resume_id}/download", headers=headers)
    assert download_resp.status_code == 200
    assert len(download_resp.content) > 0

    # Delete resume
    delete_resp = client.delete(f"/api/v1/resumes/{resume_id}", headers=headers)
    assert delete_resp.status_code == 200

    # Verify 404 after delete
    get_after_delete = client.get(f"/api/v1/resumes/{resume_id}", headers=headers)
    assert get_after_delete.status_code == 404


def test_cross_user_isolation(client: TestClient):
    user1_headers = create_auth_user(client, email="user_one@example.com")
    user2_headers = create_auth_user(client, email="user_two@example.com")
    docx_bytes = create_sample_docx_bytes()

    # User 1 uploads resume
    upload_resp = client.post(
        "/api/v1/resumes/upload",
        headers=user1_headers,
        files={"file": ("user1_cv.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    user1_resume_id = upload_resp.json()["id"]

    # User 2 attempts to get User 1's resume
    forbidden_get = client.get(f"/api/v1/resumes/{user1_resume_id}", headers=user2_headers)
    assert forbidden_get.status_code == 404

    # User 2 attempts to delete User 1's resume
    forbidden_delete = client.delete(f"/api/v1/resumes/{user1_resume_id}", headers=user2_headers)
    assert forbidden_delete.status_code == 404
