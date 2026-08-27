import io
import pytest
from docx import Document
from fastapi.testclient import TestClient


def create_auth_user(client: TestClient, email: str = "matcher_user@example.com", password: str = "SecurePass123!"):
    """Helper to register, login and return auth headers."""
    client.post(
        "/api/v1/auth/register",
        json={"email": email, "password": password, "full_name": "Matcher Candidate"},
    )
    login_resp = client.post(
        "/api/v1/auth/login",
        data={"username": email, "password": password},
    )
    token = login_resp.json()["access_token"]
    return {"Authorization": f"Bearer {token}"}


def upload_candidate_resume(client: TestClient, headers: dict) -> int:
    """Uploads a candidate resume containing Python, FastAPI, Docker skills."""
    doc = Document()
    doc.add_heading("Candidate Resume", level=1)
    doc.add_paragraph("Email: candidate@ai.com | Phone: 555-0199")
    doc.add_paragraph("Skills: Python, FastAPI, Docker, PostgreSQL, Redis, React, Git")
    doc.add_paragraph("Bachelor of Science in Computer Science, 2019")
    doc.add_paragraph("Experience: 2019 - Present (5+ years of experience) Backend Developer")
    buf = io.BytesIO()
    doc.save(buf)

    resp = client.post(
        "/api/v1/resumes/upload",
        headers=headers,
        files={"file": ("candidate_resume.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
        data={"is_primary": "true"},
    )
    return resp.json()["id"]


def test_analyze_match_with_pasted_job_description(client: TestClient):
    headers = create_auth_user(client, email="pasted_match@example.com")
    resume_id = upload_candidate_resume(client, headers)

    jd_text = """
    We are looking for a Senior Backend Developer.
    Requirements:
    - 3+ years of experience with Python, FastAPI, and Docker.
    - Experience with AWS, Kubernetes, and PostgreSQL.
    - Strong communication and unit testing practices.
    """

    payload = {
        "resume_id": resume_id,
        "company_name": "Tech Corp",
        "job_title": "Senior Backend Developer",
        "job_description": jd_text,
    }

    response = client.post("/api/v1/matching/analyze", json=payload, headers=headers)
    assert response.status_code == 201
    data = response.json()

    assert data["resume_id"] == resume_id
    assert data["company_name"] == "Tech Corp"
    assert data["match_score"] > 0
    assert "Python" in data["matching_skills"]
    assert "FastAPI" in data["matching_skills"]
    assert "AWS" in data["missing_skills"]
    assert len(data["recommendations"]) > 0


def test_analyze_match_with_tracked_job_id(client: TestClient):
    headers = create_auth_user(client, email="tracked_match@example.com")
    upload_candidate_resume(client, headers)  # uploaded as primary resume

    # Create tracked job application
    job_resp = client.post(
        "/api/v1/jobs/",
        json={
            "company_name": "Stripe",
            "job_title": "Platform Infrastructure Engineer",
            "job_description": "Seeking an engineer with Python, Docker, Kubernetes, and Redis expertise.",
            "status": "applied",
        },
        headers=headers,
    )
    job_id = job_resp.json()["id"]

    # Analyze match without explicit resume_id (falls back to primary)
    response = client.post(
        "/api/v1/matching/analyze",
        json={"job_id": job_id},
        headers=headers,
    )
    assert response.status_code == 201
    data = response.json()
    assert data["job_id"] == job_id
    assert data["company_name"] == "Stripe"
    assert "Python" in data["matching_skills"]


def test_match_history_and_get_by_id(client: TestClient):
    headers = create_auth_user(client, email="history_match@example.com")
    resume_id = upload_candidate_resume(client, headers)

    # Run two matches
    client.post(
        "/api/v1/matching/analyze",
        json={"resume_id": resume_id, "job_description": "Python, Docker engineer"},
        headers=headers,
    )
    match_resp2 = client.post(
        "/api/v1/matching/analyze",
        json={"resume_id": resume_id, "job_description": "React, TypeScript frontend"},
        headers=headers,
    )
    match_id2 = match_resp2.json()["id"]

    # History
    hist_resp = client.get("/api/v1/matching/history", headers=headers)
    assert hist_resp.status_code == 200
    assert len(hist_resp.json()) == 2

    # Get single report
    get_resp = client.get(f"/api/v1/matching/{match_id2}", headers=headers)
    assert get_resp.status_code == 200
    assert get_resp.json()["id"] == match_id2

    # Delete report
    del_resp = client.delete(f"/api/v1/matching/{match_id2}", headers=headers)
    assert del_resp.status_code == 200

    # Confirm 404
    assert client.get(f"/api/v1/matching/{match_id2}", headers=headers).status_code == 404


def test_match_cross_user_isolation(client: TestClient):
    user1_headers = create_auth_user(client, email="match_u1@example.com")
    user2_headers = create_auth_user(client, email="match_u2@example.com")
    
    upload_candidate_resume(client, user1_headers)
    
    match_resp = client.post(
        "/api/v1/matching/analyze",
        json={"job_description": "Python, FastAPI developer"},
        headers=user1_headers,
    )
    match_id = match_resp.json()["id"]

    # User 2 attempts to get User 1's report
    assert client.get(f"/api/v1/matching/{match_id}", headers=user2_headers).status_code == 404

    # User 2 attempts to delete User 1's report
    assert client.delete(f"/api/v1/matching/{match_id}", headers=user2_headers).status_code == 404
